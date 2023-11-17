import os
import shutil
import datetime
import pickle
import tkinter as tk
import psutil
from tkinter import ttk, filedialog, Toplevel, StringVar, IntVar, Label, Checkbutton, Radiobutton
from tkinter import Menu, END, DISABLED, NORMAL
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_TAB_ALIGNMENT, WD_BREAK

version = '1.1'

program_font = 'Arial'

program_dir = f'{os.getenv("APPDATA")}/LabelMaker'
instance_dir = f'{program_dir}/instance'

letters = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm',
           'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z']
capital_letters = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M',
                   'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']
numbers = ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9']
valid_chars = ['(', ')', '_', '-', ' ', *letters, *capital_letters, *numbers]


class SaveData:
    def __init__(self, label_maker: object):
        self.username_str = label_maker.username.get()
        self.address_str = label_maker.address.get()
        self.autoformat_int = label_maker.autoformat.get()
        self.food_items = label_maker.food_items
        self.groups = label_maker.groups
        self.version = label_maker.version


class FoodItem:
    def __init__(self, name: str, ingredients: str):
        self.name = name
        self.ingredients = ingredients
        self.saved_ingredients = ingredients
        self.edited = False

    def get_name(self):
        return f'*{self.name}' if self.edited else self.name

    def save_item(self):
        self.ingredients = self.ingredients.replace('\n', '')
        self.saved_ingredients = str(self.ingredients)
        self.edited = False

    def edit_item(self, ingredients: str):
        self.ingredients = ingredients.replace('\n', '')
        self.edited = (self.ingredients != self.saved_ingredients)

    def revert(self):
        self.edited = False
        self.ingredients = str(self.saved_ingredients).replace('\n', '')


class LabelGroup:
    def __init__(self, name: str, items: dict):
        self.name = name
        self.items = {**items}


class LabelMaker:
    def __init__(self):
        # Check for running instance
        running_processes = [process.name() for process in psutil.process_iter()]
        if running_processes.count('LabelMaker.exe') > 2:
            return

        # Window Properties
        self.root = tk.Tk()
        self.root.geometry('670x310')
        self.root.title('Label Maker')
        self.root.resizable(False, False)

        # Program Variables
        self.version = float(version)
        self.username = StringVar()
        self.address = StringVar()
        self.autoformat = IntVar(value=1)
        self.groups = []
        self.food_items = []
        self.food_items_dict = {}
        self.selectable_items = []
        self.selected_item = None
        self.last_used_group_name = ''
        self.label_size = (4, 1)
        self.labels_to_print = {}
        self.auto_save = False
        self.auto_save_name = ''
        self.date_stamp = datetime.datetime.now().strftime('%m-%d-%Y')
        self.save_path = f'{os.path.join(os.environ["USERPROFILE"])}/Desktop'
        self.counter = 0

        # File management
        first_run = False
        if not os.path.isdir(program_dir):
            os.mkdir(program_dir)
            first_run = True
        elif os.path.isfile(file_path := f'{program_dir}/savedata'):
            # Load saved information
            with open(file_path, 'rb') as file:
                save_data = pickle.load(file)
            if type(save_data) is not SaveData:
                raise TypeError

            try:
                if save_data.version < self.version:
                    shutil.copy(file_path, f'{file_path}-{save_data.version}')
            except AttributeError:
                shutil.copy(file_path, f'{file_path}-backup')

            load_funcs = [
                lambda: self.username.set(save_data.username_str),
                lambda: self.address.set(save_data.address_str),
                lambda: self.autoformat.set(save_data.autoformat_int),
                lambda: self.groups.extend(save_data.groups),
                lambda: self.food_items.extend(save_data.food_items),
                lambda: self.food_items_dict.update({item.name: item for item in self.food_items}),
                lambda: self.selectable_items.extend([item.get_name() for item in self.food_items])
                ]

            for func in load_funcs:
                try:
                    func()
                except AttributeError:
                    continue
        else:
            first_run = True

        # Window Menu
        self.menu = Menu(self.root)
        self.file = Menu(self.menu, tearoff=0)
        self.file.add_command(label='Settings', command=self.settings_window)
        self.menu.add_cascade(label='File', menu=self.file)
        self.root.config(menu=self.menu)

        # Program UI
        self.edit_item_label = tk.Label(self.root, text='Edit Food Item', font=(program_font, 15))
        self.edit_item_label.place(x=147, y=2, anchor='n')

        self.item_name_label = tk.Label(self.root, text='Food Item:', font=(program_font, 12))
        self.item_name_label.place(x=10, y=50, anchor='w')
        self.item_name = StringVar()
        self.item_name_box = ttk.Combobox(values=self.selectable_items, postcommand=self.dropdown_opened,
                                          textvariable=self.item_name)
        self.item_name_box.bind('<<ComboboxSelected>>', self.dropdown_changed)
        self.item_name.trace('w', self.combobox_edited)
        self.item_name_box.bind('<KeyRelease>', self.combobox_user_edit)
        self.item_name_box.place(x=90, y=50, anchor='w')

        self.edit_name_button = tk.Button(self.root, text='\u270E',
                                          command=lambda: self.edit_item_name_window(food_item=self.selected_item))
        self.edit_name_button.place(x=240, y=50, anchor='w')

        self.ingredients_label = tk.Label(self.root, text='Ingredients:', font=(program_font, 12))
        self.ingredients_label.place(x=10, y=78, anchor='w')
        self.ingredients_entry = tk.Text(self.root, width=30, height=9, font=(program_font, 12))
        self.ingredients_entry.bind('<KeyRelease>', self.textbox_edited)
        self.ingredients_entry.place(x=10, y=90, anchor='nw')

        self.save_item_button = tk.Button(self.root, text='Save Item', width=10, command=self.save_item)
        self.save_item_button.place(x=146, y=260, anchor='ne')

        self.delete_item_button = tk.Button(self.root, text='Delete Item', width=10, command=self.delete_item)
        self.delete_item_button.place(x=148, y=260, anchor='nw')

        self.separator = ttk.Separator(self.root, orient='vertical')
        self.separator.place(x=295, y=0, relwidth=0.2, relheight=1)

        self.selected_label = tk.Label(self.root, text='Labels to print:', font=(program_font, 15))
        self.selected_label.place(x=480, y=2, anchor='n')

        self.selected_item_name = StringVar()
        self.selected_label = tk.Label(self.root, textvariable=self.selected_item_name, font=(program_font, 18))
        self.selected_label.place(x=495, y=52, anchor='e')

        self.spinbox = tk.Spinbox(self.root, from_=0, to=99, width=3)
        self.spinbox.place(x=505, y=52, anchor='w')
        self.spinbox.delete(0, END)
        self.spinbox.insert(0, '1')
        self.spinbox['state'] = 'readonly'

        self.add_item_button = tk.Button(self.root, text='Add', width=3, command=self.add_item)
        self.add_item_button.place(x=540, y=52, anchor='w')

        self.clear_button = tk.Button(self.root, text='Clear', width=4, command=self.clear_labels_to_print)
        self.clear_button.place(x=615, y=95, anchor='se')

        self.items_to_print_label = tk.Label(self.root, text='Labels to print:', font=(program_font, 12))
        self.items_to_print_label.place(x=393, y=95, anchor='s')
        self.items_to_print_entry = tk.Text(self.root, width=30, height=9, font=(program_font, 12))
        self.items_to_print_entry.place(x=480, y=96, anchor='n')
        self.items_to_print_entry['state'] = DISABLED

        self.create_labels_button = tk.Button(self.root, text='Create Labels', width=10, command=self.save_labels)
        self.create_labels_button.place(x=342, y=270, anchor='nw')

        self.save_group_button = tk.Button(self.root, text='Save Group', width=10, command=self.save_label_group)
        self.save_group_button.place(x=535, y=270, anchor='ne')
        self.load_group_button = tk.Button(self.root, text='Load Group', width=10, command=self.load_label_group)
        self.load_group_button.place(x=619, y=270, anchor='ne')

        # Version Label
        self.version_label = Label(self.root, text=version)
        self.version_label.place(relx=1, rely=1.01, anchor='se')

        # Mainloop
        if first_run:
            self.root.after(1, self.on_first_run)
        self.root.protocol('WM_DELETE_WINDOW', self.on_program_exit)
        self.root.mainloop()

    def on_program_exit(self):
        if any(item.edited for item in self.food_items):
            self.save_changes()
            return
        with open(f'{program_dir}/savedata', 'wb') as file:
            save_data = SaveData(self)
            pickle.dump(save_data, file)
        self.root.destroy()

    def on_first_run(self):
        # Ask to input Username and Address
        self.settings_window(first_run=True)

    def settings_window(self, first_run=False):
        def blink_entry(entry_to_blink=None):
            if not self.counter:
                return
            if not entry_to_blink:
                self.counter = 0
                return
            if 'username' in entry_to_blink:
                if username_label_entry['background'] != 'SystemWindow':
                    username_label_entry['background'] = 'SystemWindow'
                else:
                    username_label_entry['background'] = 'pink'
                username_label_entry.update()
            if 'address' in entry_to_blink:
                if address_label_entry['background'] != 'SystemWindow':
                    address_label_entry.config(background='SystemWindow')
                    address_label_entry['background'] = 'SystemWindow'
                else:
                    address_label_entry['background'] = 'pink'
                address_label_entry.update()
            self.counter -= 1
            settings_window.after(150, blink_entry, entry_to_blink)

        def window_close():
            empty_entries = []
            if not self.username.get():
                empty_entries.append('username')
            if not self.address.get():
                empty_entries.append('address')
            if empty_entries:
                self.counter = 6
                blink_entry(empty_entries)
            else:
                settings_window.destroy()

        def autoformat_db():
            for item in self.food_items:
                item.edit_item(self.format_ingredients(item.ingredients))
                if item.name != self.format_item_name(item.name):
                    self.change_item_name(item, self.format_item_name(item.name))

        settings_window = Toplevel(self.root)
        settings_window.focus()
        settings_window.title('Edit Settings')
        settings_window.geometry('275x250')
        settings_window.geometry(f'+{self.root.winfo_rootx()}+{self.root.winfo_rooty()}')
        settings_window.resizable(False, False)
        settings_window.protocol('WM_DELETE_WINDOW', window_close)

        username_label = tk.Label(settings_window, text='Name:', font=(program_font, 12))
        username_label.place(x=80, y=15, anchor='e')
        username_label_entry = tk.Entry(settings_window, textvariable=self.username, font=(program_font, 12))
        username_label_entry.place(x=82, y=17, anchor='w')

        address_label = tk.Label(settings_window, text='Address:', font=(program_font, 12))
        address_label.place(x=80, y=45, anchor='e')
        address_label_entry = tk.Entry(settings_window, textvariable=self.address, font=(program_font, 12))
        address_label_entry.place(x=82, y=47, anchor='w')

        autoformat_check = Checkbutton(settings_window, text='Autoformat Text', variable=self.autoformat)
        autoformat_check.place(x=5, y=75)

        autoformat_button = tk.Button(settings_window, text='Autoformat Database', command=autoformat_db)
        autoformat_button.place(x=5, y=100)

        if first_run:
            settings_window.title('Enter User Information')
            settings_window.grab_set()
            settings_window.focus()
            settings_window.transient(self.root)

    def save_item(self):
        if not self.item_name_box.get():
            return
        if self.selected_item:
            ingredients = self.ingredients_entry.get('1.0', END).replace('\n', '')
            if self.autoformat:
                self.selected_item.ingredients = (ingredients := self.format_ingredients(ingredients))
                if self.selected_item.name != (new_name := self.format_item_name(self.selected_item.name)):
                    self.change_item_name(self.selected_item, new_name)
            self.selected_item.save_item()
            self.update_combobox_text_only(self.selected_item.get_name())
            self.ingredients_entry.delete('1.0', END)
            self.ingredients_entry.insert('1.0', ingredients)
        else:
            if self.autoformat:
                self.update_combobox_text_only(self.format_item_name(self.item_name_box.get()))
                ingredients = self.ingredients_entry.get('1.0', END).replace('\n', '')
                self.ingredients_entry.delete('1.0', END)
                self.ingredients_entry.insert('1.0', self.format_ingredients(ingredients))
            new_item = FoodItem(self.item_name_box.get(), self.ingredients_entry.get('1.0', END))
            self.food_items.append(new_item)
            self.food_items_dict[new_item.name] = new_item
            self.change_selected_item(new_item)
        self.update_combobox()

    def delete_item(self):
        def window_close(delete=False):
            if delete and self.selected_item:
                if self.selected_item in self.labels_to_print:
                    del self.labels_to_print[self.selected_item]
                    self.update_items_to_print_entry()
                self.selectable_items.pop(self.selectable_items.index(self.selected_item.get_name()))
                self.food_items.pop(self.food_items.index(self.selected_item))
                del self.food_items_dict[self.selected_item.name]
                self.change_selected_item()
                self.item_name_box.set('')
                self.ingredients_entry.delete('1.0', END)
                self.update_combobox()
            del_window.destroy()

        if not self.item_name_box.get():
            return

        del_window = Toplevel(self.root)
        del_window.grab_set()
        del_window.focus()
        del_window.transient(self.root)
        del_window.title('Confirm Delete')
        del_window.geometry('240x80')
        del_window.geometry(f'+{self.root.winfo_rootx()}+{self.root.winfo_rooty()}')
        del_window.resizable(False, False)
        del_window.protocol('WM_DELETE_WINDOW', window_close)

        item_name_label = tk.Label(del_window, text='Delete Item?', font=(program_font, 12))
        item_name_label.place(x=120, y=22, anchor='s')

        save_button = tk.Button(del_window, text='Delete', width=10, command=lambda: window_close(delete=True))
        save_button.place(x=115, y=32, anchor='ne')

        cancel_button = tk.Button(del_window, text='Don\'t Delete', width=10, command=window_close)
        cancel_button.place(x=125, y=32, anchor='nw')

    def change_item_name(self, food_item, new_name: str):
        if new_name in self.food_items_dict:
            del self.food_items_dict[new_name]
        del self.food_items_dict[food_item.name]
        food_item.name = new_name
        self.food_items_dict[food_item.name] = food_item
        self.update_combobox_text_only(food_item.name)
        self.update_combobox()
        self.update_items_to_print_entry()

    def edit_item_name_window(self, food_item):
        def window_close():
            def cancel_edit(*event):
                if event and event[0].keysym != 'Return':
                    return
                error_window.destroy()
                edit_name_window.destroy()

            if food_item.name != new_item_name.get():
                if new_item_name.get() in self.food_items_dict:
                    error_window = Toplevel(edit_name_window)
                    error_window.grab_set()
                    error_window.focus()
                    error_window.transient(edit_name_window)
                    error_window.title('Error')
                    error_window.geometry('225x100')
                    error_window.geometry(f'+{edit_name_window.winfo_rootx()}+{edit_name_window.winfo_rooty()}')
                    error_window.resizable(False, False)

                    existing_name_label = tk.Label(error_window, text='Name already exists', font=(program_font, 12))
                    existing_name_label.place(x=112, y=29, anchor='s')

                    edit_name_button = tk.Button(error_window, text='Cancel Edit', command=cancel_edit)
                    edit_name_button.bind('<KeyRelease>', cancel_edit)
                    edit_name_button.focus_set()
                    edit_name_button.place(x=112, y=31, anchor='n')
                    return
                self.change_item_name(food_item, new_item_name.get())
            edit_name_window.destroy()

        def key_release(event):
            if event.keysym == 'Return':
                window_close()

        if not food_item:
            return

        new_item_name = StringVar()
        new_item_name.set(food_item.name)

        edit_name_window = Toplevel(self.root)
        edit_name_window.grab_set()
        edit_name_window.focus()
        edit_name_window.transient(self.root)
        edit_name_window.title('Edit Item Name')
        edit_name_window.geometry('225x100')
        edit_name_window.geometry(f'+{self.root.winfo_rootx()}+{self.root.winfo_rooty()}')
        edit_name_window.resizable(False, False)
        edit_name_window.protocol('WM_DELETE_WINDOW', window_close)

        item_name_label = tk.Label(edit_name_window, text='Item Name:', font=(program_font, 12))
        item_name_label.place(x=112, y=29, anchor='s')
        item_name_entry = tk.Entry(edit_name_window, textvariable=new_item_name, font=(program_font, 12))
        item_name_entry.bind('<KeyRelease>', key_release)
        item_name_entry.place(x=112, y=31, anchor='n')

    @staticmethod
    def format_item_name(item_name: str):
        if len(item_name) < 2:
            return
        formatted_name = [capital_letters[letters.index(item_name[0])] if item_name[0] in letters else item_name[0]]
        capitalize = False
        for char in item_name[1:]:
            if capitalize:
                if char == ' ':
                    continue
                formatted_name.append(capital_letters[letters.index(char)] if char in letters else char)
                capitalize = False
                continue
            if char == ' ':
                capitalize = True
            formatted_name.append(char)

        return ''.join(formatted_name).rstrip(' ')

    @staticmethod
    def format_ingredients(ingredients: str):
        return ingredients.rstrip(' ') if ingredients else ''

    def save_changes(self):
        def window_close(do='cancel'):
            if do == 'cancel':
                save_window.destroy()
                return
            if do == 'save':
                for item in self.food_items:
                    item.save_item()
            else:
                for item in self.food_items:
                    item.revert()
            self.on_program_exit()

        save_window = Toplevel(self.root)
        save_window.grab_set()
        save_window.focus()
        save_window.transient(self.root)
        save_window.title('Unsaved Changes')
        save_window.geometry('240x80')
        save_window.geometry(f'+{self.root.winfo_rootx()}+{self.root.winfo_rooty()}')
        save_window.resizable(False, False)
        save_window.protocol('WM_DELETE_WINDOW', window_close)

        item_name_label = tk.Label(save_window, text='Save unsaved changes?', font=(program_font, 12))
        item_name_label.place(x=120, y=22, anchor='s')

        save_button = tk.Button(save_window, text='Save', width=10, command=lambda: window_close(do='save'))
        save_button.place(x=115, y=32, anchor='ne')

        cancel_button = tk.Button(save_window, text='Don\'t Save', width=10,
                                  command=lambda: window_close(do='dont save'))
        cancel_button.place(x=125, y=32, anchor='nw')

    def textbox_edited(self, event):
        if event.keysym in ['Right', 'Left', 'Up', 'Down']:
            return
        if event.keysym == 'Return':
            self.save_item()
            textbox_contents = self.ingredients_entry.get('1.0', END).replace('\n', '')
            self.ingredients_entry.delete('1.0', END)
            self.ingredients_entry.insert('1.0', textbox_contents)
            return
        if self.selected_item:
            self.selected_item.edit_item(self.ingredients_entry.get('1.0', END))
            self.update_combobox_text_only(self.selected_item.get_name())
        self.update_combobox()

    def combobox_edited(self, *_):
        if not self.item_name_box.get():
            return
        if self.selected_item:
            self.change_selected_item()
            self.ingredients_entry.delete('1.0', END)

    def combobox_user_edit(self, event):
        if event.keysym in ['Right', 'Left']:
            return
        value = [char for char in self.item_name_box.get() if char in valid_chars]
        value = ''.join(value).lstrip()
        self.update_combobox_text_only(value)
        if item := self.food_items_dict.get(value):
            self.ingredients_entry.delete('1.0', END)
            self.ingredients_entry.insert('1.0', item.ingredients.replace('\n', ''))
            self.change_selected_item(item)
            self.update_combobox_text_only(item.get_name())

    def dropdown_changed(self, *_):
        if not (item_name := self.item_name_box.get().replace('*', '')):
            return
        if self.auto_save and self.auto_save_name != self.item_name_box.get():
            if self.autoformat:
                self.auto_save_name = self.format_item_name(self.auto_save_name)
            new_item = FoodItem(self.auto_save_name, self.ingredients_entry.get('1.0', END))
            new_item.edited = True
            self.food_items.append(new_item)
            self.food_items_dict[new_item.name] = new_item
            self.update_combobox()
            self.auto_save_name = ''
        self.auto_save = False
        if item := self.food_items_dict.get(item_name):
            self.ingredients_entry.delete('1.0', END)
            self.ingredients_entry.insert('1.0', item.ingredients.replace('\n', ''))
            self.change_selected_item(item)

    def dropdown_opened(self, *_):
        if not self.selected_item and self.item_name_box.get():
            self.auto_save = True
            self.auto_save_name = self.item_name_box.get()

    def update_combobox(self):
        self.selectable_items = [item.get_name() for item in self.food_items]
        self.item_name_box['values'] = self.selectable_items

    def update_combobox_text_only(self, value: str):
        selected_item = self.selected_item
        textbox_contents = self.ingredients_entry.get('1.0', END).replace('\n', '')
        self.item_name.set(value)
        self.ingredients_entry.delete('1.0', END)
        self.ingredients_entry.insert('1.0', textbox_contents)
        self.change_selected_item(selected_item)

    def add_item(self):
        if not self.selected_item:
            return
        if self.spinbox.get() == '0' and self.labels_to_print.get(self.selected_item):
            del self.labels_to_print[self.selected_item]
        elif self.spinbox.get() == '0':
            return
        else:
            self.labels_to_print[self.selected_item] = self.spinbox.get()
        self.update_items_to_print_entry()

    def clear_labels_to_print(self):
        if not self.labels_to_print:
            return

        def clear_entry(do_clear=True):
            if do_clear:
                self.labels_to_print = {}
                self.update_items_to_print_entry()
            confirm_clear_pop_up.destroy()

        confirm_clear_pop_up = Toplevel(self.root)
        confirm_clear_pop_up.title('Confirm Clear')
        confirm_clear_pop_up.geometry('240x70')
        confirm_clear_pop_up.geometry(f'+{self.root.winfo_rootx() + 370}+{self.root.winfo_rooty()}')
        confirm_clear_pop_up.grab_set()
        confirm_clear_pop_up.focus()
        confirm_clear_pop_up.transient(self.root)
        confirm_clear_pop_up.resizable(False, False)

        confirm_label = Label(confirm_clear_pop_up, text='Clear labels to print?')
        confirm_label.place(x=120, y=10, anchor='n')

        yes_button = tk.Button(confirm_clear_pop_up, text='Yes', width='10', command=clear_entry)
        yes_button.place(x=115, y=65, anchor='se')

        no_button = tk.Button(confirm_clear_pop_up, text='No', width='10', command=lambda: clear_entry(do_clear=False))
        no_button.place(x=125, y=65, anchor='sw')

    def update_items_to_print_entry(self):
        labels_to_print = []
        for item in self.labels_to_print:
            labels_to_print.extend([f'{item.name} ({self.labels_to_print[item]})', '\n'])
        labels_to_print = ''.join(labels_to_print[:-1])
        self.items_to_print_entry['state'] = NORMAL
        self.items_to_print_entry.delete('1.0', END)
        self.items_to_print_entry.insert('1.0', labels_to_print)
        self.items_to_print_entry['state'] = DISABLED

    def create_labels(self):
        if not self.labels_to_print:
            return None

        # Format Document
        document = Document()
        document.styles['Normal'].paragraph_format.space_before = Pt(0)
        document.styles['Normal'].paragraph_format.space_after = Pt(0)
        document.styles['Normal'].font.name = 'Arial'
        for section in document.sections:
            section.page_width = Inches(self.label_size[0])
            section.page_height = Inches(self.label_size[1])
            margin_size = Inches(1/16)
            section.top_margin = margin_size
            section.bottom_margin = margin_size
            section.left_margin = margin_size
            section.right_margin = margin_size

        # Create document contents
        for i, item in enumerate(self.labels_to_print):
            for j in range(limit := int(self.labels_to_print[item])):
                para1 = document.add_paragraph()
                item_name = para1.add_run(f'{item.name}')
                item_name.bold = True
                item_name.font.size = Pt(11)
                para1.add_run(f'\t{self.username.get()} - {self.address.get()}').font.size = Pt(10)
                margin_end = Inches((sec := document.sections[0]).page_width.inches -
                                    (sec.left_margin.inches + sec.right_margin.inches))
                tab_stops = para1.paragraph_format.tab_stops
                tab_stops.add_tab_stop(margin_end, WD_TAB_ALIGNMENT.RIGHT)

                para2 = document.add_paragraph()
                ingredients = para2.add_run('Ingredients: ')
                ingredients.bold = True
                ingredients.font.size = Pt(9)
                para2.add_run(item.ingredients.replace('\n', '')).font.size = Pt(8)
                para3 = document.add_paragraph()
                date = para3.add_run('Date: ')
                date.font.size = Pt(8)
                date_stamp = para3.add_run('')  # Automated date stamp goes here
                date_stamp.font.size = Pt(8)
                date_stamp.bold = True
                weight_spacing = ''.join([' ' for _ in range(40)])
                weight = para3.add_run(f'{weight_spacing}Weight:')
                weight.font.size = Pt(8)
                if i == len(self.labels_to_print) - 1 and j == limit - 1:
                    continue
                weight.add_break(WD_BREAK.PAGE)

        return document

    def save_labels(self):
        if not self.labels_to_print:
            return
        out_file = filedialog.asksaveasfile(initialdir=self.save_path, initialfile=f'Labels {self.date_stamp}.docx',
                                            filetypes=[('Word Document', '*.docx')])
        if not out_file:
            return
        if not (out_file_path := out_file.name).endswith('.docx'):
            out_file.close()
            os.rename(out_file_path, (out_file_path := f'{out_file_path}.docx'))
        self.save_path = os.path.dirname(out_file_path)
        document = self.create_labels()
        document.save(out_file_path)

    def save_label_group(self, *_):
        if not self.labels_to_print:
            return

        if not self.last_used_group_name:
            group_name = ''
            for group in self.groups:
                if group.items == self.labels_to_print:
                    group_name = group.name
                    break
        else:
            group_name = self.last_used_group_name

        def save_and_close():
            def confirm_overwrite():
                def do_overwrite():
                    self.groups.pop(index_of_group := self.groups.index(existing_group))
                    self.groups.insert(index_of_group, LabelGroup(new_group_name_str, self.labels_to_print))
                    self.last_used_group_name = new_group_name_str
                    confirm_overwrite_pop_up.destroy()
                    save_group_window.destroy()

                confirm_overwrite_pop_up = Toplevel(save_group_window)
                confirm_overwrite_pop_up.title('Overwrite Existing Group?')
                confirm_overwrite_pop_up.geometry('240x70')
                confirm_overwrite_pop_up.geometry(
                    f'+{save_group_window.winfo_rootx()}+{save_group_window.winfo_rooty()}')
                confirm_overwrite_pop_up.grab_set()
                confirm_overwrite_pop_up.focus()
                confirm_overwrite_pop_up.transient(save_group_window)
                confirm_overwrite_pop_up.resizable(False, False)

                confirm_label = Label(confirm_overwrite_pop_up, text=f'Overwrite "{new_group_name_str}" group?')
                confirm_label.place(x=120, y=10, anchor='n')

                yes_button = tk.Button(confirm_overwrite_pop_up, text='Yes', width='10', command=do_overwrite)
                yes_button.place(x=115, y=65, anchor='se')

                no_button = tk.Button(confirm_overwrite_pop_up, text='No', width='10',
                                      command=confirm_overwrite_pop_up.destroy)
                no_button.place(x=125, y=65, anchor='sw')
            new_group_name_str = new_group_name.get()
            existing_group = None
            for group_ in self.groups:
                if group_.name == new_group_name_str:
                    existing_group = group_
                    break
            if existing_group:
                confirm_overwrite()
                return
            self.groups.append(LabelGroup(new_group_name_str, self.labels_to_print))
            self.last_used_group_name = new_group_name_str
            save_group_window.destroy()

        def key_release(event):
            if event.keysym == 'Return':
                save_and_close()

        new_group_name = StringVar()
        new_group_name.set(group_name)

        save_group_window = Toplevel(self.root)
        save_group_window.grab_set()
        save_group_window.focus()
        save_group_window.transient(self.root)
        window_title = 'Save Group'
        save_group_window.title(window_title)
        save_group_window.geometry('225x100')
        save_group_window.geometry(f'+{self.root.winfo_rootx()}+{self.root.winfo_rooty()}')
        save_group_window.resizable(False, False)

        group_name_label = tk.Label(save_group_window, text='Group Name:', font=(program_font, 12))
        group_name_label.place(x=112, y=29, anchor='s')
        group_name_entry = tk.Entry(save_group_window, textvariable=new_group_name, font=(program_font, 12))
        group_name_entry.bind('<KeyRelease>', key_release)
        group_name_entry.place(x=112, y=31, anchor='n')

        save_group_button = tk.Button(save_group_window, text='Save', width=10, command=save_and_close)
        save_group_button.place(x=107, y=60, anchor='ne')

        cancel_button = tk.Button(save_group_window, text='Cancel', width=10, command=save_group_window.destroy)
        cancel_button.place(x=117, y=60, anchor='nw')

    def load_label_group(self, *_):
        if not self.groups:
            return

        def load_group():
            if selected_radio.get() < 0:
                return

            self.labels_to_print = {**radio_group_dict[selected_radio.get()].items}
            self.update_items_to_print_entry()
            self.last_used_group_name = radio_group_dict[selected_radio.get()].name
            load_group_window.destroy()

        def confirm_delete():
            if selected_radio.get() < 0:
                return

            def delete_group(do_delete=True):
                if do_delete:
                    self.groups.pop(self.groups.index(radio_group_dict[selected_radio.get()]))
                    radio_button = radio_buttons[selected_radio.get()]
                    radio_button.config(text=f'{radio_button.cget("text")} - DELETED')
                    selected_radio.set(-1)
                    radio_button['state'] = DISABLED
                confirm_del_pop_up.destroy()

            confirm_del_pop_up = Toplevel(load_group_window)
            confirm_del_pop_up.title('Confirm Delete')
            confirm_del_pop_up.geometry('240x70')
            confirm_del_pop_up.geometry(f'+{load_group_window.winfo_rootx()}+{load_group_window.winfo_rooty()}')
            confirm_del_pop_up.grab_set()
            confirm_del_pop_up.focus()
            confirm_del_pop_up.transient(load_group_window)
            confirm_del_pop_up.resizable(False, False)

            confirm_label = Label(confirm_del_pop_up,
                                  text=f'Delete "{radio_group_dict[selected_radio.get()].name}" group?')
            confirm_label.place(x=120, y=10, anchor='n')

            yes_button = tk.Button(confirm_del_pop_up, text='Yes', width='10', command=delete_group)
            yes_button.place(x=115, y=65, anchor='se')

            no_button = tk.Button(confirm_del_pop_up, text='No', width='10',
                                  command=lambda: delete_group(do_delete=False))
            no_button.place(x=125, y=65, anchor='sw')

        load_group_window = Toplevel(self.root)
        load_group_window.grab_set()
        load_group_window.focus()
        load_group_window.transient(self.root)
        window_title = 'Select Group'
        load_group_window.title(window_title)
        win_height, win_width = 300, 275
        load_group_window.geometry(f'{win_height}x{win_width}')
        load_group_window.geometry(f'+{self.root.winfo_rootx()}+{self.root.winfo_rooty()}')
        load_group_window.resizable(False, False)

        load_group_button = tk.Button(load_group_window, text='Load', width=10, command=load_group)
        load_group_button.place(x=(win_width / 2) - 5, y=win_height - 28, anchor='se')

        delete_group_button = tk.Button(load_group_window, text='Delete', width=10, command=confirm_delete)
        delete_group_button.place(x=(win_width / 2) + 5, y=win_height - 28, anchor='sw')

        container = ttk.Frame(load_group_window)
        canvas = tk.Canvas(container, height=win_height - 58, width=win_width - 12)
        scrollbar = ttk.Scrollbar(container, orient='vertical', command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        scrollable_frame.bind('<Configure>', lambda _: canvas.configure(scrollregion=canvas.bbox('all')))
        canvas.create_window((0, 0), window=scrollable_frame, anchor='nw')
        canvas.configure(yscrollcommand=scrollbar.set)
        if len(self.groups) > 10:
            canvas.bind_all('<MouseWheel>',
                            lambda event: canvas.yview_scroll(int(-1 * (event.delta / 120)), 'units'))
        else:
            canvas.bind_all('<MouseWheel>', lambda _: None)

        radio_group_dict = {}
        selected_radio = IntVar()
        radio_buttons = {}
        for i, group in enumerate(self.groups):
            radio_group_dict[i] = group
            radio_buttons[i] = Radiobutton(scrollable_frame, text=group.name, variable=selected_radio, value=i)
            radio_buttons[i].grid(sticky='w', row=i, column=0)
        selected_radio.set(-1)

        container.pack()
        canvas.pack(side='left', fill='both', expand=True)
        if len(self.groups) > 10:
            scrollbar.pack(side='right', fill='y')

    def change_selected_item(self, item=None):
        item_name = item.name if item else ''
        self.selected_item = item
        self.selected_item_name.set(item_name)
        spinbox_text = '1'
        if self.selected_item in self.labels_to_print:
            spinbox_text = self.labels_to_print[self.selected_item]
        self.spinbox['state'] = NORMAL
        self.spinbox.delete(0, END)
        self.spinbox.insert(0, spinbox_text)
        self.spinbox['state'] = 'readonly'


if __name__ == '__main__':
    LabelMaker()
